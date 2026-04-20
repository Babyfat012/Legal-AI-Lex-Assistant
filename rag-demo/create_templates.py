"""
Script tạo DOCX templates cho Document Generation.

Chạy script này 1 lần để tạo các file .docx template trong thư mục templates/.
Mỗi template sử dụng Jinja2 tags ({{ variable }}) để docxtpl render.
"""

import os
import sys

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)


def _set_style(doc):
    """Set default font for the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)


def _add_header(doc, quoc_hieu=True):
    """Add standard Vietnamese document header."""
    if quoc_hieu:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        run.bold = True
        run.font.size = Pt(13)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Độc lập - Tự do - Hạnh phúc")
        run2.bold = True
        run2.font.size = Pt(13)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run("---o0o---")
        run3.font.size = Pt(13)


def create_don_ly_hon():
    """Tạo template Đơn xin ly hôn."""
    doc = Document()
    _set_style(doc)
    _add_header(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐƠN XIN LY HÔN")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # Kính gửi
    p = doc.add_paragraph()
    run = p.add_run("Kính gửi: ")
    run.bold = True
    p.add_run("{{ toa_an }}")

    doc.add_paragraph()

    # Người yêu cầu
    p = doc.add_paragraph()
    run = p.add_run("NGƯỜI YÊU CẦU:")
    run.bold = True

    doc.add_paragraph("Họ và tên: {{ ho_ten_nguoi_yeu_cau }}")
    doc.add_paragraph("Ngày sinh: {{ ngay_sinh_nguoi_yeu_cau }}")
    doc.add_paragraph("CCCD/CMND số: {{ cccd_nguoi_yeu_cau }}")
    doc.add_paragraph("Địa chỉ thường trú: {{ dia_chi_nguoi_yeu_cau }}")
    doc.add_paragraph("Số điện thoại: {{ sdt_nguoi_yeu_cau }}")

    doc.add_paragraph()

    # Bên kia
    p = doc.add_paragraph()
    run = p.add_run("VỢ/CHỒNG:")
    run.bold = True

    doc.add_paragraph("Họ và tên: {{ ho_ten_ben_kia }}")
    doc.add_paragraph("Ngày sinh: {{ ngay_sinh_ben_kia }}")
    doc.add_paragraph("CCCD/CMND số: {{ cccd_ben_kia }}")
    doc.add_paragraph("Địa chỉ thường trú: {{ dia_chi_ben_kia }}")

    doc.add_paragraph()

    # Nội dung
    p = doc.add_paragraph()
    run = p.add_run("NỘI DUNG YÊU CẦU:")
    run.bold = True

    doc.add_paragraph("Ngày đăng ký kết hôn: {{ ngay_ket_hon }}")
    doc.add_paragraph("Nơi đăng ký kết hôn: {{ noi_ket_hon }}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Lý do ly hôn:")
    run.bold = True
    doc.add_paragraph("{{ ly_do }}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Về con chung:")
    run.bold = True
    doc.add_paragraph("{{ con_chung }}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Về tài sản chung:")
    run.bold = True
    doc.add_paragraph("{{ tai_san }}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Tôi xin cam đoan những lời khai trên là đúng sự thực "
        "và xin chịu trách nhiệm trước pháp luật."
    )

    doc.add_paragraph()

    # Signature
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("........., ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p2.add_run("Người yêu cầu")
    run.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run("(Ký và ghi rõ họ tên)")

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.add_run("{{ ho_ten_nguoi_yeu_cau }}")

    filepath = os.path.join(TEMPLATES_DIR, "don_ly_hon.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


def create_don_khieu_nai():
    """Tạo template Đơn khiếu nại."""
    doc = Document()
    _set_style(doc)
    _add_header(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐƠN KHIẾU NẠI")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Kính gửi: ")
    run.bold = True
    p.add_run("{{ co_quan }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("NGƯỜI KHIẾU NẠI:")
    run.bold = True

    doc.add_paragraph("Họ và tên: {{ ho_ten }}")
    doc.add_paragraph("Ngày sinh: {{ ngay_sinh }}")
    doc.add_paragraph("CCCD/CMND số: {{ cccd }}")
    doc.add_paragraph("Địa chỉ thường trú: {{ dia_chi }}")
    doc.add_paragraph("Số điện thoại: {{ sdt }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("ĐỐI TƯỢNG KHIẾU NẠI:")
    run.bold = True
    doc.add_paragraph("{{ doi_tuong_khieu_nai }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("NỘI DUNG KHIẾU NẠI:")
    run.bold = True
    doc.add_paragraph("{{ noi_dung }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("YÊU CẦU GIẢI QUYẾT:")
    run.bold = True
    doc.add_paragraph("{{ yeu_cau }}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Tôi xin cam đoan những nội dung khiếu nại trên là đúng sự thật "
        "và xin chịu trách nhiệm trước pháp luật về nội dung khiếu nại."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("........., ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p2.add_run("Người khiếu nại")
    run.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run("(Ký và ghi rõ họ tên)")

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.add_run("{{ ho_ten }}")

    filepath = os.path.join(TEMPLATES_DIR, "don_khieu_nai.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


def create_don_khoi_kien():
    """Tạo template Đơn khởi kiện."""
    doc = Document()
    _set_style(doc)
    _add_header(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐƠN KHỞI KIỆN")
    run.bold = True
    run.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("(V/v: Khởi kiện vụ án dân sự)")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Kính gửi: ")
    run.bold = True
    p.add_run("{{ toa_an }}")

    doc.add_paragraph()

    # Nguyên đơn
    p = doc.add_paragraph()
    run = p.add_run("NGUYÊN ĐƠN:")
    run.bold = True

    doc.add_paragraph("Họ và tên: {{ ho_ten_nguyen_don }}")
    doc.add_paragraph("Ngày sinh: {{ ngay_sinh_nguyen_don }}")
    doc.add_paragraph("CCCD/CMND số: {{ cccd_nguyen_don }}")
    doc.add_paragraph("Địa chỉ thường trú: {{ dia_chi_nguyen_don }}")
    doc.add_paragraph("Số điện thoại: {{ sdt_nguyen_don }}")

    doc.add_paragraph()

    # Bị đơn
    p = doc.add_paragraph()
    run = p.add_run("BỊ ĐƠN:")
    run.bold = True

    doc.add_paragraph("Họ và tên: {{ ho_ten_bi_don }}")
    doc.add_paragraph("Địa chỉ: {{ dia_chi_bi_don }}")

    doc.add_paragraph()

    # Nội dung
    p = doc.add_paragraph()
    run = p.add_run("NỘI DUNG KHỞI KIỆN:")
    run.bold = True
    doc.add_paragraph("{{ noi_dung_khoi_kien }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("YÊU CẦU TÒA ÁN GIẢI QUYẾT:")
    run.bold = True
    doc.add_paragraph("{{ yeu_cau_khoi_kien }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("TÀI LIỆU, CHỨNG CỨ KÈM THEO:")
    run.bold = True
    doc.add_paragraph("{{ chung_cu }}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Tôi xin cam đoan những lời khai trên là đúng sự thực, "
        "nếu sai tôi xin chịu trách nhiệm trước pháp luật."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("........., ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p2.add_run("Người khởi kiện")
    run.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run("(Ký và ghi rõ họ tên)")

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.add_run("{{ ho_ten_nguyen_don }}")

    filepath = os.path.join(TEMPLATES_DIR, "don_khoi_kien.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")


if __name__ == "__main__":
    create_don_ly_hon()
    create_don_khieu_nai()
    create_don_khoi_kien()
    print("All templates created successfully!")
